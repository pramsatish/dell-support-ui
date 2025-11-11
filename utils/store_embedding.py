import os
from docx import Document as DocxDocument
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
import re

# -----------------------------
# Configuration
# -----------------------------
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
DOCS_DIR = os.path.join(BASE_DIR, "docs", "dell-data")
CHROMA_DB_DIR = os.path.join(BASE_DIR, "chroma_dell_db")
EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"  # Hugging Face model

# -----------------------------
# Folder setup
# -----------------------------
def ensure_directories():
    """Ensure required folders exist, or create them."""
    os.makedirs(DOCS_DIR, exist_ok=True)
    os.makedirs(CHROMA_DB_DIR, exist_ok=True)
    print(f"ğŸ“ Verified folders: {DOCS_DIR} and {CHROMA_DB_DIR}")

# -----------------------------
# Document Reading
# -----------------------------
def read_docx(filepath: str) -> str:
    """Extracts and returns text content from a .docx file."""
    doc = DocxDocument(filepath)
    text_blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(text_blocks)

# -----------------------------
# Semantic Chunking
# -----------------------------
def semantic_chunk_text(text):
    text = re.sub(r'\n{3,}', '\n\n', text.strip())
    chunks = re.split(r'\n\s*\n', text)
    chunks = [c.strip() for c in chunks if len(c.strip()) > 0]
    return chunks

# -----------------------------
# Document Preparation
# -----------------------------
def process_all_docs(docs_folder: str):
    """Parses all .docx files, splits into semantic chunks, returns LangChain Document objects."""
    documents = []
    for file in sorted(os.listdir(docs_folder)):
        if file.lower().endswith(".docx"):
            filepath = os.path.join(docs_folder, file)
            print(f"ğŸ“„ Processing file: {file}")

            text = read_docx(filepath)
            if not text:
                print(f"âš ï¸ Skipped empty file: {file}")
                continue

            chunks = semantic_chunk_text(text)
            for idx, chunk in enumerate(chunks):
                metadata = {"source": file, "chunk": idx}
                documents.append(Document(page_content=chunk, metadata=metadata))

            print(f"âœ… {len(chunks)} semantic chunks created for {file}")
    print(f"\nğŸ“š Total Chunks Ready: {len(documents)}")
    return documents

# -----------------------------
# Embedding + ChromaDB Storage
# -----------------------------
def create_and_store_embeddings(documents):
    """Generates embeddings using Hugging Face and stores them in ChromaDB."""
    print("\nğŸš€ Creating embeddings using Hugging Face model...")
    embeddings = HuggingFaceEmbeddings(model_name=EMBED_MODEL)

    vectordb = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        persist_directory=CHROMA_DB_DIR
    )

    #vectordb.persist()
    print(f"âœ… Embeddings stored in: {CHROMA_DB_DIR}")

# -----------------------------
# Main Entry Point
# -----------------------------
if __name__ == "__main__":
    ensure_directories()
    docs = process_all_docs(DOCS_DIR)
    if docs:
        create_and_store_embeddings(docs)
        print("\nğŸ‰ Embedding creation completed successfully for Dell Knowledge Base!")
    else:
        print("âŒ No .docx files found in the docs/dell-data folder. Add files and re-run.")

from dotenv import load_dotenv
from pathlib import Path
import os
import sys
import re
from docx import Document as DocxDocument
from chromadb import Client
from chromadb.utils import embedding_functions
from groq import Groq
import google.generativeai as genai  

#Load environment variables
env_path = Path(__file__).resolve().parent.parent / ".env"
print(" Loading .env from:", env_path)

load_dotenv(dotenv_path=env_path)

# Optional check 
#print(" GEMINI_API_KEY:", os.getenv("GEMINI_API_KEY"))
#print(" GROQ_API_KEY:", os.getenv("GROQ_API_KEY"))


# Disable tokenizers parallelism warnings
os.environ["TOKENIZERS_PARALLELISM"] = "false"


# Configuration
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
DOCS_DIR = os.path.join(BASE_DIR, "docs", "dell-data")
MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"


# API Keys
GROQ_KEY = os.getenv("GROQ_API_KEY")
GEMINI_KEY = os.getenv("GEMINI_API_KEY")

if not GROQ_KEY and not GEMINI_KEY:
    print(" No LLM API keys found. Set GROQ_API_KEY or GEMINI_API_KEY in environment.")
    sys.exit(1)


# Initialize LLM Clients
try:
    groq_client = Groq(api_key=GROQ_KEY) if GROQ_KEY else None
except Exception as e:
    groq_client = None
    print(f" Groq client failed to initialize: {e}")

try:
    if GEMINI_KEY:
        genai.configure(api_key=GEMINI_KEY)
        gemini_model = genai.GenerativeModel("gemini-2.5-pro")
    else:
        gemini_model = None
except Exception as e:
    gemini_model = None
    print(f" Gemini client failed to initialize: {e}")


# Document & Chunking Utilities
class Document:
    def __init__(self, page_content, metadata=None):
        self.page_content = page_content
        self.metadata = metadata or {}


def read_docx(filepath):
    doc = DocxDocument(filepath)
    return "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])


def semantic_chunk_text(text):
    text = re.sub(r'\n{3,}', '\n\n', text.strip())
    chunks = re.split(r'\n\s*\n', text)
    return [c.strip() for c in chunks if c.strip()]


# Chroma Vectorstore
def create_vectorstore():
    if not os.path.exists(DOCS_DIR) or not os.listdir(DOCS_DIR):
        print(f" No .docx files found in {DOCS_DIR}")
        sys.exit(1)

    docs = []
    for file in sorted(os.listdir(DOCS_DIR)):
        if file.lower().endswith(".docx"):
            filepath = os.path.join(DOCS_DIR, file)
            text = read_docx(filepath)
            if not text:
                continue
            chunks = semantic_chunk_text(text)
            for idx, chunk in enumerate(chunks):
                docs.append(Document(page_content=chunk, metadata={"source": file, "chunk": idx}))
            print(f" {len(chunks)} semantic chunks created from {file}")

    embedding_func = embedding_functions.SentenceTransformerEmbeddingFunction(model_name=MODEL_NAME)
    chroma = Client()
    vectordb = chroma.create_collection(name="dell_kb", embedding_function=embedding_func)
    vectordb.add(
        documents=[d.page_content for d in docs],
        metadatas=[d.metadata for d in docs],
        ids=[str(i) for i in range(len(docs))],
    )
    print(f"\n Vectorstore created with {len(docs)} total semantic chunks\n")
    return vectordb


def load_vectorstore():
    chroma = Client()
    try:
        return chroma.get_collection(name="dell_kb")
    except Exception:
        print("ℹ Chroma DB not found. Creating from docs...")
        return create_vectorstore()


# AI Query with Fallback
def get_ai_response(query, docs):
    context = "\n\n".join(docs)
    prompt = f"""
You are a Dell technical support assistant.
Use the following Dell documentation to answer the question.

Context:
{context}

Question: {query}

Answer clearly and helpfully.
"""

    # Primary: Groq
    if groq_client:
        try:
            response = groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(f" Groq failed: {e}")

    # Fallback: Gemini
    if gemini_model:
        try:
            gemini_response = gemini_model.generate_content(prompt)
            return gemini_response.text.strip()
        except Exception as e2:
            print(f" Gemini failed: {e2}")

    return " No LLM available to generate a response."


# Interactive Query Loop
def interactive_query(vectordb):
    print("\n Dell Knowledge Base \n")
    while True:
        query = input("Q: ").strip()
        if not query or query.lower() in ("exit", "quit"):
            print(" Bye!")
            break

        results = vectordb.query(query_texts=[query], n_results=4)
        docs = results["documents"][0] if results else []

        if not docs:
            print("No results found.\n")
            continue

        print("\n Top Matching Documents:")
        for i, doc in enumerate(docs, 1):
            snippet = doc[:200].replace("\n", " ")
            print(f" {i}. {snippet}...\n")

        print(" Generating AI response...\n")
        ai_answer = get_ai_response(query, docs)
        print(" AI Answer:\n")
        print(ai_answer)
        print("\n" + "-" * 60 + "\n")


# Run App
if __name__ == "__main__":
    vectordb = load_vectorstore()
    interactive_query(vectordb)
